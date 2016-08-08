#!/usr/bin/env python
# coding:utf-8

import paramiko
from multiprocessing import Process, Queue
import time
import os
from docx import Document
from openpyxl import Workbook
from functools import wraps
import numpy as np
from conf import CONF
import logging


def wait_time(times, info):
    def decorate(func):
        @wraps(func)
        def wrapper(self, *args, **kwargs):
            self.start_time = time.time()
            self.output_q.put('{} {} start time:'.format(self.pair_info['{role}_manage_ip'.format(role=self.role)], info) + str(self.start_time))
            func(self, *args, **kwargs)
            self.end_time = time.time()
            self.output_q.put('{} {} end time:'.format(self.pair_info['{role}_manage_ip'.format(role=self.role)], info) + str(self.end_time))
            while True:
                self.end_time = time.time()
                self.use_time = self.end_time - self.start_time
                if self.use_time < times:
                    time.sleep(1)
                else:
                    break
            self.output_q.put('{} {} use time:'.format(self.pair_info['{role}_manage_ip'.format(role=self.role)], info) + str(self.use_time))
            return func
        return wrapper
    return decorate


class HostProcess(Process):
    def __init__(self, pair_info, mtu, output_q, role, cmd_q, result_q, s_check_q=None):
        Process.__init__(self)
        self.pair_info = pair_info
        self.mtu = mtu
        self.output_q = output_q
        self.role = role
        self.conn = None
        self.cmd_q = cmd_q
        self.result_q = result_q
        self.s_check_q = s_check_q

    @wait_time(10, 'connect')
    def con_ssh(self):
        self.conn = paramiko.SSHClient()
        self.conn.set_missing_host_key_policy(paramiko.AutoAddPolicy())
        while True:
            try:
                self.conn.connect(self.pair_info['{role}_manage_ip'.format(role=self.role)],
                                  self.pair_info['{role}_ssh_port'.format(role=self.role)],
                                  self.pair_info['{role}_ssh_user'.format(role=self.role)],
                                  self.pair_info['{role}_ssh_password'.format(role=self.role)])
                self.output_q.put('connect')
                break
            except Exception, e:
                self.output_q.put(e)

    def change_mtu(self, device):
        getmtu = 'cat /sys/class/net/{}/mtu'.format(device)
        changemtu = "echo '{}'>/sys/class/net/{}/mtu".format(self.mtu, device)
        while True:
            stdin, stdout, stderr = self.conn.exec_command(getmtu)
            cmd_result = stdout.read(), stderr.read()
            cmd_result = ''.join(cmd_result).strip()
            if cmd_result != self.mtu:
                self.conn.exec_command(changemtu)
                self.output_q.put(u'{} {} change mtu {}--->{}'.format(self.pair_info['{role}_manage_ip'.format(role=self.role)], device, cmd_result, self.mtu))
            else:
                self.output_q.put(u'{} {} current mtu {}'.format(self.pair_info['{role}_manage_ip'.format(role=self.role)], device, cmd_result))
                break

    def run_cmd(self, cmd):
        stdin, stdout, stderr = self.conn.exec_command(cmd)
        if self.s_check_q is not None:
            self.s_check_q.put('server')
        self.output_q.put('{} run cmd {} at {}'.format(self.pair_info['{role}_manage_ip'.format(role=self.role)], cmd, str(time.time())))
        cmd_result = stdout.read(), stderr.read()
        result = ''.join(cmd_result).strip()
        self.result_q.put(result)
        self.output_q.put('{} end cmd {} at {}'.format(self.pair_info['{role}_manage_ip'.format(role=self.role)], cmd, str(time.time())))

    def run(self):
        if self.conn is None:
            self.con_ssh()
        if self.pair_info['{role}_change_mtu'.format(role=self.role)]:
            for device in self.pair_info['{role}_change_mtu_device'.format(role=self.role)]:
                self.change_mtu(device)
            time.sleep(30)
        while True:
            time.sleep(2)
            cmd = self.cmd_q.get()
            if cmd == 'over':
                self.conn.close()
                break
            self.run_cmd(cmd)


class OutProcess(Process):
    def __init__(self, output_q, loger):
        Process.__init__(self)
        self.output_q = output_q
        self.loger = loger

    def run(self):
        while True:
            out = self.output_q.get()
            self.loger.debug(out)
            if out != 'over':
                print out
            else:
                break


class WriteProcess(Process):
    def __init__(self, loger, num, collect_all_result, scenes_dir, pair, node_mtu, all_cmd_q, pair_num):
        Process.__init__(self)
        self.loger = loger
        self.num = num
        self.collect_all_result = collect_all_result
        self.pair = pair
        self.scenes_dir = scenes_dir
        self.node_mtu = node_mtu
        self.output_q = Queue()
        self.s_cmd_q = Queue()
        self.c_cmd_q = Queue()
        self.s_result = Queue()
        self.c_result = Queue()
        self.s_check_q = Queue()
        self.all_cmd_q = all_cmd_q
        self.pair_num = pair_num

    def run(self):
        for vm_mtu in CONF.vm_mtu_list:
            out_put = OutProcess(self.output_q, self.loger)
            out_put.start()
            self.writeresult(vm_mtu)
            out_put.join()

    def writeresult(self, vm_mtu):
        ori_path, extract_path = self.create_file(vm_mtu)
        ori_docx = Document()
        extract_xlsx = Workbook()
        self.create_excel_sheet(extract_xlsx)
        column = 1
        self.write_ori_head(ori_docx)
        s = HostProcess(CONF[self.pair], vm_mtu, self.output_q, 'server', self.s_cmd_q, self.s_result, self.s_check_q)
        c = HostProcess(CONF[self.pair], vm_mtu, self.output_q, 'client', self.c_cmd_q, self.c_result)
        s.start()
        c.start()
        run_cmd_times = 0
        for item in CONF.cmd_list:
            ori_docx.add_paragraph(item.upper(), style='Heading 4')
            s_cmd = self.create_cmd('server', item)
            c_cmd = self.create_cmd('client', item)
            for seq in range(CONF.test_times):
                self.s_cmd_q.put(s_cmd)
                while not self.s_cmd_q.empty():
                    time.sleep(0.2)
                self.s_check_q.get()
                time.sleep(2)
                self.c_cmd_q.put(c_cmd)
                try:
                    s_result = self.s_result.get(timeout=240)
                    c_result = self.c_result.get(timeout=240)
                except:
                    self.kill_process(CONF[self.pair], 'server')
                    s_result = self.s_result.get(timeout=240)
                    c_result = self.c_result.get(timeout=240)
                self.output_q.put('wait all cmd run complete')
                self.write_ori_to_docx(ori_docx, seq, s_cmd, c_cmd, s_result, c_result)
                extracted_result = self.extract_result(c_result)
                ws = extract_xlsx.get_sheet_by_name(u'第{}组数据'.format(str(seq+1)))
                for index, result_cell in enumerate(extracted_result):
                    self.output_q.put(extracted_result)
                    self.output_q.put(result_cell)
                    ws.cell(row=2, column=column+index, value=result_cell)
                    all_result = self.collect_all_result.get()
                    all_result[seq][self.num][column+index] = result_cell
                    self.collect_all_result.put(all_result)
                run_cmd_times += 1
                self.output_q.put(run_cmd_times)
                self.all_cmd_q.put('cmd_run')
                while self.all_cmd_q.qsize() != run_cmd_times * self.pair_num:
                    time.sleep(0.3)
                time.sleep(3)
            column += len(extracted_result)
        ori_docx.save(ori_path)
        extract_xlsx.save(extract_path)
        self.output_q.put('save_xlsx')
        self.s_cmd_q.put('over')
        self.c_cmd_q.put('over')
        s.join()
        c.join()
        self.output_q.put('over')

    def kill_process(self, pair_info, role):
        conn = paramiko.SSHClient()
        conn.set_missing_host_key_policy(paramiko.AutoAddPolicy())
        while True:
            try:
                conn.connect(pair_info['{}_manage_ip'.format(role)],
                             pair_info['{}_ssh_port'.format(role)],
                             pair_info['{}_ssh_user'.format(role)],
                             pair_info['{}_ssh_password'.format(role)])
                break
            except Exception, e:
                self.output_q.put(e)
        getpid = 'ps aux|grep iperf3|grep -v grep'
        killpid = "kill {}"
        stdin, stdout, stderr = conn.exec_command(getpid)
        cmd_result = stdout.read(), stderr.read()
        pid = cmd_result[0].split()[1]
        conn.exec_command(killpid.format(pid))

    @staticmethod
    def write_ori_to_docx(ori_docx, seq, s_cmd, c_cmd, s_result, c_result):
            ori_docx.add_paragraph(u'第{}组数据'.format(str(seq+1)), style='Heading 5')
            table = ori_docx.add_table(rows=0, cols=1, style='Normal Table')
            row_cells = table.add_row().cells
            row_cells[0].text = 'client'
            row_cells = table.add_row().cells
            row_cells[0].text = c_cmd
            result_list = c_result.split('\n')
            if len(result_list) > 50:
                result_list = result_list[:20]+result_list[-20:]
                result_list = (r for r in result_list)
            for result in result_list:
                time.sleep(0.1)
                row_cells = table.add_row().cells
                row_cells[0].text = result
            ori_docx.add_paragraph(u'\n', style='Normal')

            table = ori_docx.add_table(rows=0, cols=1, style='Normal Table')
            row_cells = table.add_row().cells
            row_cells[0].text = 'server'
            row_cells = table.add_row().cells
            row_cells[0].text = s_cmd
            result_list = s_result.split('\n')
            if len(result_list) > 50:
                result_list = result_list[:20]+result_list[-20:]
                result_list = (r for r in result_list)
            for result in result_list:
                time.sleep(0.1)
                row_cells = table.add_row().cells
                row_cells[0].text = result
            ori_docx.add_paragraph(u'\n', style='Normal')

    @staticmethod
    def write_cmd_to_docx(cmd_docx, s_cmd, c_cmd):
        table = cmd_docx.add_table(rows=0, cols=1, style='Normal Table')
        row_cells = table.add_row().cells
        row_cells[0].text = 'client'
        row_cells = table.add_row().cells
        row_cells[0].text = c_cmd
        cmd_docx.add_paragraph(u'\n', style='Normal')

        row_cells = table.add_row().cells
        row_cells[0].text = 'server'
        row_cells = table.add_row().cells
        row_cells[0].text = s_cmd
        cmd_docx.add_paragraph(u'\n', style='Normal')

    @staticmethod
    def extract_result(c_result):
        result_list = (r for r in c_result.split('\n')[-5:])
        for result in result_list:
            time.sleep(0.2)
            if '0.00-{}.00'.format(CONF.iperf_time) and 'receiver' in result:
                a = result.split()[6]
                if 'Mb' in result.split()[7]:
                    a = float(result.split()[6]) / 1000
                return (a,)
            elif '0.00-{}.00'.format(CONF.iperf_time) and '%)' in result:
                if 'Mb' in result.split()[7]:
                    bw = float(result.split()[6]) / 1000
                else:
                    bw = result.split()[6]
                delay = result.split()[8]
                lose = result.split()[11].strip(')').strip('(').strip('%')
                capacity = float(bw)*(100-float(lose))/100
                return bw, delay, float(lose), capacity
            elif 'error' in result:
                space = 0
                return space, space, space, space

    def write_ori_head(self, ori_docx):
        ori_docx.add_paragraph(u'测试目的', style='Heading 1')
        ori_docx.add_paragraph(u'此处填写测试目的。', style='Normal')
        ori_docx.add_paragraph(u'测试环境', style='Heading 1')
        ori_docx.add_paragraph(u'此处填写测试环境。', style='Normal')
        ori_docx.add_paragraph(u'环境信息', style='Heading 2')
        ori_docx.add_paragraph(u'此处填写环境信息', style='Normal')
        ori_docx.add_paragraph(u'在物理网卡上配置IP', style='Heading 2')
        ori_docx.add_paragraph(u'MTU {}'.format(self.node_mtu), style='Heading 3')

    @staticmethod
    def create_excel_sheet(wb):
        for i in range(CONF.test_times):
            a = wb.create_sheet(title=u'第{}组数据'.format(str(i+1)))
            head = []
            for cmd in CONF.cmd_list:
                if 'TCP' in cmd:
                    head.append(u'吞吐Gb({})'.format(cmd))
                elif 'UDP' in cmd:
                    head.append(u'带宽Gb({})'.format(cmd))
                    head.append(u'时延ms({})'.format(cmd))
                    head.append(u'丢包率%({})'.format(cmd))
                    head.append(u'吞吐Gb({})'.format(cmd))
            a.append(head)

    def create_cmd(self, role, item):
        command = CONF['{role}_cmd'.format(role=role)][item].format(server_data_ip=CONF[self.pair]['server_manage_ip'],
                                                                    iperf_time=CONF['iperf_time'])
        return command

    def get_dir(self):
        dir_name = os.path.join(self.scenes_dir, self.pair)
        if not os.path.exists(dir_name):
            os.makedirs(dir_name)
        return dir_name

    @staticmethod
    def get_file_path(file_name, dir_name, file_list):
        if file_name in file_list:
            if file_name.endswith('docx'):
                file_path = os.path.join(dir_name, file_name.replace('.docx', 'I.docx'))
            else:
                file_path = os.path.join(dir_name, file_name.replace('.xlsx', 'I.xlsx'))
        else:
            file_path = os.path.join(dir_name, file_name)
        return file_path

    def create_file(self, mtu):
        dir_name = self.get_dir()
        # 定义文件名
        ori_filename = '{}.docx'.format(CONF.filename.format(vmmtu=mtu, nodemtu=self.node_mtu, pair_num=self.pair_num))
        extract_filename = '{}.xlsx'.format(CONF.filename.format(vmmtu=mtu, nodemtu=self.node_mtu, pair_num=self.pair_num))
        file_list = os.listdir(dir_name)
        ori_path = self.get_file_path(ori_filename, dir_name, file_list)
        extract_path = self.get_file_path(extract_filename, dir_name, file_list)
        return ori_path, extract_path
